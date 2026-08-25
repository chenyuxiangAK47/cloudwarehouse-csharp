# -*- coding: utf-8 -*-
"""Clean AI placeholder chrome in Final-Report-EN - 副本.docx (in place)."""
from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DOC = Path(r"d:\tools\cloudwarehouse-csharp\docs\project-management\Final-Report-EN - 副本.docx")


def has_drawing(para) -> bool:
    for el in para._p.iter():
        if el.tag.endswith("}drawing") or el.tag.endswith("}pict"):
            return True
    return False


def clear_text_keep_runs_structure(para) -> None:
    """Remove all text from paragraph but keep drawing runs if any."""
    # Safer: if has drawing, only clear text nodes in non-drawing runs
    if has_drawing(para):
        for run in para.runs:
            # skip runs that contain drawings
            if any(c.tag.endswith("}drawing") or c.tag.endswith("}pict") for c in run._r.iter()):
                continue
            run.text = ""
        return
    para.clear()


def set_para_text(para, text: str, *, bold: bool = False) -> None:
    """Replace paragraph text; drop paragraph if empty and no drawing."""
    if has_drawing(para):
        # keep image; strip instructional text only
        for run in para.runs:
            if any(c.tag.endswith("}drawing") or c.tag.endswith("}pict") for c in run._r.iter()):
                continue
            run.text = ""
        return
    # clear and write one run
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    run = para.add_run(text)
    run.bold = bold


def extract_caption(block: str) -> str | None:
    m = re.search(r"\*图注：([^*]+)\*", block)
    if m:
        return m.group(1).strip()
    m = re.search(r"图注[:：]\s*(.+)", block)
    if m:
        return m.group(1).strip().rstrip("*").strip()
    return None


def extract_fid(block: str) -> str | None:
    m = re.search(r"【插图占位\s*([^】]+)】", block)
    return m.group(1).strip() if m else None


CAPTION_EN = {
    "1-1": "CloudWarehouse admin entry — runnable system evidence.",
    "3-1": "CloudWarehouse actors and use cases; PDA use cases in §3.6.",
    "3-2": "Waybill dual-track preview: machine vs sheet amounts.",
    "3-3": "PDA no-order report success — closed-loop evidence.",
    "4-1": "Project roadmap milestones (Phase 1 + Phase 2).",
    "4-2": "Phase 1 solo Planned vs Actual hours.",
    "5-1": "Entity-relationship diagram (schema.sql authoritative).",
    "6-1": "Logical architecture — layers and module dependencies.",
    "6-2": "DDD bounded contexts (Master Data / Import / Pricing / …).",
    "6-3": "Enterprise context map — CloudWarehouse and PDA independent; integration Planned.",
    "6-4": "Physical / deployment topology — single-instance; no HA.",
    "7-1": "Billing Strategy class diagram (Tier / Overweight / Volumetric).",
    "7-2": "Waybill dual-track preview sequence.",
    "8-1": "CI pipeline activity (CI delivered; full CD not claimed).",
    "8-2": "GitHub Actions successful run.",
    "8-3": "Coverage Summary from CI artefact (no hard-coded % slogan).",
    "9-1": "Risk register overview (project / technical / security).",
}


def main() -> None:
    assert DOC.exists(), DOC
    backup = DOC.with_name(
        DOC.stem + f".bak-{datetime.now().strftime('%Y%m%d-%H%M%S')}" + DOC.suffix
    )
    shutil.copy2(DOC, backup)
    print("backup:", backup)

    doc = Document(str(DOC))

    # Title page paragraphs 0-2 typically
    # Scan and fix by content
    removed_paste = 0
    cleaned_fig = 0
    for para in doc.paragraphs:
        t = para.text
        ts = t.strip()

        if "最终实习报告（中文整理稿）" in t or "Chinese Master Draft" in t:
            set_para_text(para, "Final Internship Report", bold=True)
            continue

        if "插图占位已标出" in t or ("Solo Intern" in t and "NUS MTech SE33" in t and "粘贴" in t):
            set_para_text(para, "")
            # empty title subtitle — leave blank or remove content
            continue

        if ts.startswith("▼▼▼") or "在下方空白处粘贴图片" in t:
            if has_drawing(para):
                clear_text_keep_runs_structure(para)
            else:
                set_para_text(para, "")
            removed_paste += 1
            continue

        if "【插图占位" in t:
            fid = extract_fid(t) or "?"
            cap = CAPTION_EN.get(fid) or extract_caption(t) or fid
            set_para_text(para, f"Figure {fid}  {cap}", bold=True)
            cleaned_fig += 1
            continue

        if "*(Figure placeholder" in t or "Figure placeholder — paste PNG" in t:
            set_para_text(para, "")
            continue

        if "11.4 Client Feedback (Placeholder)" in t:
            set_para_text(para, "11.4 Client Feedback", bold=True)
            continue

        if "【文字占位】" in t:
            # keep Chinese note for senior as requested earlier
            set_para_text(
                para,
                "（待企业导师补充演示反馈摘要：日期、主要意见、已闭环项。）",
            )
            continue

        # strip emoji from paragraph text
        if any(s in t for s in ("✅", "🔄", "❌", "⚠️")):
            if has_drawing(para):
                continue
            newt = (
                t.replace("✅ ", "")
                .replace("✅", "")
                .replace("🔄 ", "")
                .replace("🔄", "")
                .replace("❌ ", "")
                .replace("❌", "")
                .replace("⚠️ ", "")
                .replace("⚠️", "")
            )
            if newt != t:
                set_para_text(para, newt)

    # Tables: emoji status
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    if not any(s in t for s in ("✅", "🔄", "❌", "⚠️")):
                        continue
                    newt = (
                        t.replace("✅ ", "")
                        .replace("✅", "")
                        .replace("🔄 ", "")
                        .replace("🔄", "")
                        .replace("❌ ", "")
                        .replace("❌", "")
                        .replace("⚠️ ", "")
                        .replace("⚠️", "")
                    )
                    if newt != t:
                        # rewrite cell paragraph carefully
                        for r in list(p.runs):
                            r._r.getparent().remove(r._r)
                        p.add_run(newt)

    # Soften English intro line if it screams AI/master md
    for para in doc.paragraphs[:15]:
        if "English counterpart of `Final-Report-ZH-Master.md`" in para.text:
            set_para_text(
                para,
                "This document is the English final internship report for the CloudWarehouse "
                "freight settlement system and the PDA no-order reporting application.",
            )

    doc.save(str(DOC))
    print("saved:", DOC)
    print("figures cleaned:", cleaned_fig, "paste markers removed:", removed_paste)
    print("size:", DOC.stat().st_size)


if __name__ == "__main__":
    main()
