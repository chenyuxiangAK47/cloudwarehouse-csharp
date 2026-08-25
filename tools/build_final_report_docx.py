# -*- coding: utf-8 -*-
"""Assemble Chinese final report → Markdown + DOCX (clean report style, no AI paste chrome)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DRAFT = ROOT / "草稿文件"
CH6 = ROOT / "docs/project-management/final-report-ch6-architecture-zh.md"
OUT_MD = ROOT / "docs/project-management/Final-Report-ZH-Master.md"
OUT_DOCX = ROOT / "docs/project-management/Final-Report-ZH.docx"
OUT_DOCX_ALT = ROOT / "docs/project-management/Final-Report-ZH-updated.docx"

FIGURE_CSS = (
    "\n\n**图 {fid}** {caption}\n\n"
)

CH11 = """
# 第十一章 结论与展望

### 11.1 结论

本实习在 Solo 条件下交付了两套并列系统：CloudWarehouse（Modular Monolith 运费结算 MVP，含 Phase 2 的 Strategy 计费、应收/应付双轨与历史价、**内置规则 RAG** 辅助查阅）与霍尼韦尔 PDA 无订单报工 MVP。二者服务同一工厂目标，但按限界上下文独立演进，**本期未做生产级 API 打通**。

中期反馈已通过可验证产物回应：多视角架构图与诚实无 HA 声明、Strategy 类图与双轨时序、CI/CodeQL/测试证据、个人 Planned vs Actual（Phase 1：198→211 小时）。规则 RAG 仅作 FAQ 检索增强，**不替代** FeeCalculationEngine。

### 11.2 已知限制

- 无 JWT/RBAC；CORS/HTTP 为演示配置
- 无生产级 HA / 完整 CD / DAST 常态门禁
- 体积重引擎已通，运单 Excel 主路径仍以实重为主
- 异形件/罚款等计费变体仍为 Planned
- 规则 RAG 为词法检索（非生产级向量语义 RAG）；未配置 ApiKey 时为摘录生成

### 11.3 展望（量化方向）

| 项 | 状态 | 依赖 |
|----|------|------|
| JWT + RBAC | Planned | ADR |
| 演示环境 DAST 基线 | Planned | 稳定演示部署 |
| 完整 CD | Planned | 认证与发布目标环境 |
| 微服务提取 | Planned | 触发条件（见第六章） |
| 云仓↔PDA 集成 | Planned | 稳定 ID/文件交换约定 |

### 11.4 Client Feedback

（待企业导师补充演示反馈摘要：日期、主要意见、已闭环项。）

### 11.5 提交物

- 本报告（中文定稿）
- 英文版
- 评估演示视频
- 附录证据截图（见附录 A）
"""

APPENDIX = """
# 附录 A 证据与截图清单

附录证据一览：

| 编号 | 内容 | 建议来源 |
|------|------|----------|
| A-01 | GitHub Actions CI 绿勾 | Actions 网页 |
| A-02 | coverage Summary | CI Artifact |
| A-03 | CodeQL 成功 | Actions |
| A-04 | 价表导入成功 | 管理端 UI |
| A-05 | 价表导入失败/未入库 | 管理端 UI |
| A-06 | 运单双轨预览一致/不一致 | 管理端 UI |
| A-07 | 工时柱状图 | sprint-hours-chart.html |
| A-08 | PDA 开工/报工 | 设备或模拟器 |
| A-09 | 规则 RAG 查询结果（含三步流水线） | 管理端「规则 RAG」Tab |
| A-10 | 非法扩展名上传被拒（可选） | UI |
| A-11 | 解决方案结构 / Modules 目录（可选） | IDE |

# 附录 B 术语与禁话速查

| 正确 | 禁止 |
|------|------|
| 双轨=应收报价 vs 应付成本 | 国内/国际线路 |
| Modular Monolith；无 HA | 微服务已上线；生产多活已建成 |
| Strategy Tier/Overweight/Volumetric Done | JSON 规则引擎；AI 智能计费 |
| 内置规则 RAG（词法检索 FAQ） | AI 结算 / 向量 RAG 已生产上线 |
| PDA 未与云仓结算 API 打通 | 已打通 / Parallel Data Aggregator |
| 覆盖率以 Artifact 为准 | 正文写死 >80% |
"""


def fig(fid: str, title: str, src: str, caption: str) -> str:
    return FIGURE_CSS.format(fid=fid, title=title, src=src, caption=caption)


def fix_text(s: str) -> str:
    reps = [
        ("智能导入", "结构化 Excel 导入（自动探测表头）"),
        ("计价规则检索功能，提供知识库辅助查阅，不作为结算引擎",
         "内置规则 RAG（Retrieve→Augment→Generate）辅助查阅业务规则 FAQ，不作为结算引擎"),
        ("计价规则检索", "内置规则 RAG"),
        ("规则检索前端页面", "规则 RAG 前端页面（流水线可视化）"),
        ("Rule knowledge lookup", "Built-in Rule RAG"),
        ("规则知识检索（辅助，非结算真相源）", "内置规则 RAG（辅助 FAQ，非结算真相源）"),
        ("通过知识库查询获取规则解释（只读）", "内置规则 RAG：检索知识库并生成带引用回答（只读）"),
        ("规则搜索界面", "规则 RAG 界面"),
        ("Assistant模块仅支持信息检索", "Assistant/规则 RAG 模块仅支持查阅检索"),
        ("规则检索 Assistant（仅用于查阅，不参与金额结算）",
         "内置规则 RAG / Assistant（仅查阅，不参与金额结算）"),
        ("Assistant 规则检索", "内置规则 RAG"),
        ("仓库`log`目录", "仓库根目录 `log` 文件"),
        ("仓库`log` 目录", "仓库根目录 `log` 文件"),
        ("留存于仓库`log`目录", "留存于仓库根目录 `log` 文件"),
        ("产线/班组/设备", "产线/机群/机床"),
        ("选择产线/班组/设备", "选择产线/机群/机床"),
        ("docs/project management/", "docs/project-management/"),
        ("sprint hours chart data.csv", "sprint-hours-chart-data.csv"),
        ("sprint hours chart.html", "sprint-hours-chart.html"),
        ("10 roadmap milestones.puml", "10-roadmap-milestones.puml"),
        ("16-context-map.puml", "16-enterprise-context-map.puml"),
        ("前： 已登录。", "前： 系统可访问（MVP 认证延期）。"),
        ("Database Design and Multi-View Architecture（接第五章）", "System Architecture and Multi-View Design"),
        ("生成 BillLineTotals 汇总对象", "调用 BillLineTotals 进行汇总与对比"),
        (
            "针对同一条运输车道（lane）且同一生效日期的规则，执行“先删除后插入”的 upsert 逻辑，整批次覆盖该车道下的全部规则行",
            "针对同一运输车道（SiteId + DestId）执行“先删除该 lane 全部规则再插入”的替换逻辑（非整生效日局部 upsert）",
        ),
        (
            "针对“站点+目的地+生效日期”的核心查询路径建立组合索引，大幅提升价格规则匹配效率。配套脚本 `fix-price-rules-index.sql` 核心作用为**删除错误的唯一索引约束**——因单车道下存在一对多多档位规则，唯一索引与业务模型冲突，移除后通过组合索引保障查询性能。",
            "配套脚本 `database/fix-price-rules-index.sql` 的核心作用是**删除错误的唯一索引** `(SiteId, DestId, EffectiveDate)`——因同一 lane、同一生效日必须允许多条档位/续重行。当前查询以非唯一的 lane 索引（如 SiteId+DestId）为主，与一对多基数对齐。",
        ),
        ("Sprint 1 4", "Sprint 1–4"),
        ("S4 终期", "S4–终期"),
        ("预览 确认", "预览–确认"),
        ("Sprint3 Sprint4", "Sprint 3–4"),
        ("Dual track", "Dual-track"),
        ("No order", "No-order"),
        ("辅助功能，用于查阅而非结算", "内置规则 RAG：查阅 FAQ，不参与结算"),
    ]
    for a, b in reps:
        s = s.replace(a, b)
    return s


def convert_tab_tables(text: str) -> str:
    """Convert blocks after '表格' with tab-separated lines into Markdown tables."""
    lines = text.splitlines()
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "表格" or line.strip().startswith("表格"):
            i += 1
            rows: list[list[str]] = []
            while i < len(lines):
                raw = lines[i]
                if not raw.strip():
                    break
                if raw.startswith("#") or re.match(r"^\d+\.\d+", raw) or raw.startswith("本章") or raw.startswith("单人") or raw.startswith("图示") or raw.startswith("工时分析") or raw.startswith("迭代") or raw.startswith("偏差") or raw.startswith("超支") or raw.startswith("改进") or raw.startswith("定稿") or raw.startswith("Phase") or raw.startswith("整体") or raw.startswith("跨平台") or raw.startswith("表述") or raw.startswith("当前策略") or raw.startswith("以上控制") or raw.startswith("需要明确") or raw.startswith("该流水线") or raw.startswith("环境适配") or raw.startswith("测试体系") or raw.startswith("SAST ") or raw.startswith("正文禁止") or raw.startswith("值得") or raw.startswith("所有列出") or raw.startswith("根据") or raw.startswith("必须") or raw.startswith("应该") or raw.startswith("可以") or raw.startswith("不会") or raw.startswith("MoSCoW") or raw.startswith("主要用例") or raw.startswith("本节") or raw.startswith("模块：") or raw.startswith("用例：") or raw.startswith("表 ") or raw.startswith("图 ") or raw.startswith("CloudWarehouse") or raw.startswith("PDA ") or raw.startswith("在 ") or raw.startswith("需特别") or raw.startswith("下表") or raw.startswith("以下") or raw.startswith("按请求") or raw.startswith("双栈") or raw.startswith("可插") or raw.startswith("•") or raw.startswith("1.") or raw.startswith("2.") or raw.startswith("3.") or raw.startswith("4.") or raw.startswith("5.") or raw.startswith("6.") or raw.startswith("7.") or raw.startswith("8.") or raw.startswith("o\t") or raw.startswith("**【") or raw.startswith("链路") or raw.startswith("Phase 1") or raw.startswith("该设计") or raw.startswith("策略模式") or raw.startswith("CreateDefault") or raw.startswith("解析顺序") or raw.startswith("所有策略") or raw.startswith("体积重") or raw.startswith("新增计费") or raw.startswith("当前处于") or raw.startswith("运单双轨") or raw.startswith("双轨语义") or raw.startswith("确认入库") or raw.startswith("历史价") or raw.startswith("成本价表") or raw.startswith("从软件") or raw.startswith("本章以") or raw.startswith("第七章") or raw.startswith("第八章") or raw.startswith("针对") or raw.startswith("本章质量") or raw.startswith("流水线") or raw.startswith("核心工作") or raw.startswith("触发规则") or raw.startswith("运行环境") or raw.startswith("标准执行") or raw.startswith("代码覆盖") or raw.startswith("报告附录") or raw.startswith("对应工作") or raw.startswith("CI 流水线") or raw.startswith("两类环境") or raw.startswith("本项目测试") or raw.startswith("本章性能") or raw.startswith("当前系统在"):
                    break
                if "\t" in raw:
                    rows.append([c.strip() for c in raw.split("\t")])
                    i += 1
                else:
                    break
            if rows:
                out.append("")
                header = rows[0]
                out.append("| " + " | ".join(header) + " |")
                out.append("| " + " | ".join(["---"] * len(header)) + " |")
                for r in rows[1:]:
                    # pad
                    while len(r) < len(header):
                        r.append("")
                    out.append("| " + " | ".join(r[: len(header)]) + " |")
                out.append("")
            continue
        out.append(line)
        i += 1
    return "\n".join(out)


def inject_figures(md: str) -> str:
    injections = [
        (
            "1.7 交付物快照",
            fig("1-1", "系统首页 / 管理端总览（可选）", "本机运行截图 wwwroot/index.html", "CloudWarehouse 管理端入口，证明可运行系统。"),
        ),
        (
            "3.5 用例图",
            fig("3-1", "用例图 Use Case Diagram", "docs/diagrams/06-use-case-diagram.puml", "CloudWarehouse 与外部参与者关系；PDA 用例见正文表。"),
        ),
        (
            "一张来自PDA显示成功报工的截图（TODO: 截图）。",
            "一张来自PDA显示成功报工的截图。"
            + fig("3-2", "运单双轨预览 UI", "管理端运单预览截图", "应收/应付机器值与表内值对比。")
            + fig("3-3", "PDA 报工成功", "PDA 设备或模拟器截图", "无订单报工闭环证据。"),
        ),
        (
            "4.2 项目里程碑总览",
            fig("4-1", "项目路线图里程碑", "docs/diagrams/10-roadmap-milestones.puml", "若图中 M6 仍为 Planned，以正文 Done 为准。"),
        ),
        (
            "图表建议：截取",
            fig("4-2", "Phase 1 个人工时 Planned vs Actual", "docs/project-management/sprint-hours-chart.html", "Solo 个人工时柱状图；数据见 sprint-hours-chart-data.csv。")
            + "图表建议（若上图已贴可删本句）：截取",
        ),
        (
            "5.4 实体关系图（ERD）",
            fig("5-1", "实体关系图 ERD", "docs/diagrams/07-erd.puml", "以 schema.sql 为准；图若滞后于 BillLines/CustomerQuoteRules 请在 caption 说明。"),
        ),
        (
            "6.2 逻辑架构与典型请求流",
            fig("6-1", "逻辑架构图", "docs/diagrams/02-logical-architecture.puml", "分层与模块依赖。"),
        ),
        (
            "6.3 限界上下文与代码映射",
            fig("6-2", "DDD 限界上下文", "docs/diagrams/05-ddd-bounded-contexts.puml", "Master Data / Import / Pricing 等边界。"),
        ),
        (
            "6.4 企业级上下文关系（含 PDA）",
            fig("6-3", "企业 Context Map", "docs/diagrams/16-enterprise-context-map.puml", "云仓与 PDA 独立；集成为 Planned。"),
        ),
        (
            "6.5 物理部署与运行拓扑",
            fig("6-4", "物理 / 部署视图", "docs/diagrams/03-physical-architecture.puml 或 04-deployment-diagram.puml", "单实例拓扑；无 HA。"),
        ),
        (
            "7.3 Strategy 类设计",
            fig("7-1", "计费 Strategy 类图", "docs/diagrams/13-billing-strategy-class.puml", "Tier / Overweight / Volumetric + FeeCalculationEngine。"),
        ),
        (
            "7.5 运单双轨时序（详细设计）",
            fig("7-2", "运单双轨预览时序图", "docs/diagrams/14-sequence-waybill-dual-track.puml", "应收 CustomerQuoteRules vs 应付 PriceRules。"),
        ),
        (
            "流水线活动图对应文件：docs/diagrams/09-cicd-pipeline.puml。",
            "流水线活动图对应文件：`docs/diagrams/09-cicd-pipeline.puml`。"
            + fig("8-1", "CI/CD 活动图", "docs/diagrams/09-cicd-pipeline.puml", "CI 为主；完整 CD 未宣称。")
            + fig("8-2", "GitHub Actions 成功运行", "Actions 网页截图", "绿勾证据。")
            + fig("8-3", "覆盖率 Summary", "CI Artifact coverage-report", "勿在正文写死百分比口号。"),
        ),
        (
            "风险治理流程图示文件为",
            fig("9-1", "风险登记示意", "docs/diagrams/12-risk-management.puml", "项目/技术/安全三类风险。")
            + "风险治理流程图示文件为",
        ),
    ]
    for anchor, block in injections:
        if anchor in md and block not in md:
            md = md.replace(anchor, anchor + "\n" + block, 1)
    return md


def dedupe_and_insert_ch6(raw: str) -> str:
    # Remove first duplicated Chapter 1 (lines before second "第一章 项目概述")
    parts = raw.split("第一章 项目概述")
    if len(parts) >= 3:
        raw = "第一章 项目概述" + parts[2]
    elif len(parts) == 2:
        raw = "第一章 项目概述" + parts[1]

    # Draft often lacks explicit "第三章" title
    ch3_anchor = "本章在第二章技术架构的基础上"
    if "第三章" not in raw.split(ch3_anchor)[0][-200:] and ch3_anchor in raw:
        raw = raw.replace(
            ch3_anchor,
            "第三章 系统用例与业务模块\nSystem Use Cases and Business Modules\n\n" + ch3_anchor,
            1,
        )

    ch6_text = CH6.read_text(encoding="utf-8")
    ch6_text = fix_text(ch6_text)
    if not ch6_text.startswith("#"):
        ch6_text = "# " + ch6_text

    marker = "第七章 软件设计"
    if marker in raw and "第六章 系统架构设计" not in raw.split(marker)[0]:
        raw = raw.replace(marker, "\n\n" + ch6_text + "\n\n# " + marker, 1)
    return raw


def normalize_headings(md: str) -> str:
    chapter_re = re.compile(
        r"^第[一二三四五六七八九十百零〇两]+章\s+.{1,40}$"
    )
    bad_chapter = re.compile(r"阐述|界定|完整呈现|从多视角|聚焦|说明上述|系统说明")
    lines = []
    for line in md.splitlines():
        s = line.strip()
        if chapter_re.match(s) and not bad_chapter.search(s) and not s.startswith("#"):
            line = "# " + s
        elif re.match(r"^\d+\.\d+\.\d+\s", s) and not s.startswith("#"):
            line = "### " + s
        elif re.match(r"^\d+\.\d+\s", s) and not s.startswith("#"):
            line = "## " + s
        lines.append(line)
    return "\n".join(lines)


def ensure_ch10_checklist(md: str) -> str:
    if "10.7" not in md and "## 10.8 本章小结" in md:
        block = """
## 10.7 仍待附录补齐的截图清单（作者执行）

定稿前建议逐项核对：

- [ ] GitHub Actions CI 绿勾
- [ ] 覆盖率 Summary
- [ ] CodeQL 成功
- [ ] 价表导入成功 / 失败
- [ ] 运单双轨预览
- [ ] 工时柱状图
- [ ] PDA 开工/报工
- [ ]（可选）非法扩展名被拒

"""
        md = md.replace("## 10.8 本章小结", block + "## 10.8 本章小结")
    return md


def set_run_font(run, size=11, bold=None, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "宋体")


def md_to_docx(md: str, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CloudWarehouse 云仓运费结算与 PDA 无订单报工")
    r.bold = True
    r.font.size = Pt(18)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("最终实习报告")
    r2.font.size = Pt(14)
    doc.add_page_break()

    # Skip duplicate H1 title already used on cover when present in md
    lines = md.splitlines()
    if lines and lines[0].startswith("# CloudWarehouse"):
        # drop title + optional **最终实习报告** + blank
        lines = lines[1:]
        while lines and (not lines[0].strip() or lines[0].strip() == "**最终实习报告**"):
            lines = lines[1:]

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Markdown table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?---", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^:?-+:?$", row[0].replace(" ", "")):
                    rows.append(row)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell = table.rows[ri].cells[ci]
                        cell.text = row[ci] if ci < len(row) else ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                set_run_font(run, size=9, bold=(ri == 0))
                doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("> "):
            para = doc.add_paragraph(stripped[2:])
            if para.runs:
                para.runs[0].italic = True
        elif stripped.startswith("- [ ]") or stripped.startswith("- [x]"):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("- ") or stripped.startswith("• ") or stripped.startswith("•\t"):
            text = re.sub(r"^[•\-]\s*", "", stripped)
            doc.add_paragraph(text, style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped) or re.match(r"^[o]\s", stripped):
            text = re.sub(r"^[o\d]+\.\s*", "", stripped)
            text = re.sub(r"^o\s+", "", text)
            doc.add_paragraph(text, style="List Number")
        elif stripped.startswith("```"):
            i += 1
            continue
        else:
            # bold **x**
            para = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    set_run_font(run, size=11, bold=True)
                else:
                    run = para.add_run(part)
                    set_run_font(run, size=11)
        i += 1

    doc.save(path)


def export_cleaned_master() -> None:
    """Rebuild DOCX from already-cleaned Final-Report-ZH-Master.md (preferred)."""
    md = OUT_MD.read_text(encoding="utf-8")
    try:
        md_to_docx(md, OUT_DOCX)
        print(f"Wrote {OUT_DOCX}")
    except PermissionError:
        md_to_docx(md, OUT_DOCX_ALT)
        print(f"Original DOCX locked; wrote {OUT_DOCX_ALT}")


def main() -> None:
    import sys

    if "--from-master" in sys.argv:
        export_cleaned_master()
        return

    raw = DRAFT.read_text(encoding="utf-8")
    raw = dedupe_and_insert_ch6(raw)
    raw = fix_text(raw)
    raw = convert_tab_tables(raw)
    raw = normalize_headings(raw)
    raw = ensure_ch10_checklist(raw)
    raw = inject_figures(raw)

    front = (
        "# CloudWarehouse 云仓运费结算与 PDA 无订单报工\n\n"
        "**最终实习报告**\n\n"
    )
    md = front + raw + "\n" + CH11 + "\n" + APPENDIX
    md = fix_text(md)
    # strip any residual AI figure chrome if inject used old cache
    md = re.sub(r"> \*\*【插图占位.*?(?=\n(?!>)|\Z)", "", md, flags=re.S)

    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(md, encoding="utf-8")
    export_cleaned_master()
    print(f"Wrote {OUT_MD}")
    print(f"MD chars: {len(md)}")


if __name__ == "__main__":
    main()
